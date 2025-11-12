"""
Export utilities for different output formats.
"""

import json
from pathlib import Path
from typing import Dict, List

from docx import Document
from docx.shared import Pt

from util.logging import get_logger

logger = get_logger(__name__)


class Exporter:
    """Export OCR results to various formats."""

    @staticmethod
    def export_txt(text: str, output_path: Path):
        """
        Export as plain text.

        Args:
            text: Text to export
            output_path: Output file path
        """
        try:
            output_path.parent.mkdir(parents=True, exist_ok=True)
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text)
            logger.info(f"Exported TXT to {output_path}")
        except Exception as e:
            logger.error(f"Failed to export TXT: {e}")
            raise

    @staticmethod
    def export_json(data: Dict, output_path: Path, pretty: bool = True):
        """
        Export as JSON.

        Args:
            data: Data to export
            output_path: Output file path
            pretty: Pretty print JSON
        """
        try:
            output_path.parent.mkdir(parents=True, exist_ok=True)
            with open(output_path, 'w', encoding='utf-8') as f:
                if pretty:
                    json.dump(data, f, indent=2, ensure_ascii=False)
                else:
                    json.dump(data, f, ensure_ascii=False)
            logger.info(f"Exported JSON to {output_path}")
        except Exception as e:
            logger.error(f"Failed to export JSON: {e}")
            raise

    @staticmethod
    def export_markdown(text: str, output_path: Path, title: str = None):
        """
        Export as Markdown.

        Args:
            text: Text to export
            output_path: Output file path
            title: Optional document title
        """
        try:
            output_path.parent.mkdir(parents=True, exist_ok=True)

            content = []
            if title:
                content.append(f"# {title}\n\n")

            content.append(text)

            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(''.join(content))

            logger.info(f"Exported Markdown to {output_path}")
        except Exception as e:
            logger.error(f"Failed to export Markdown: {e}")
            raise

    @staticmethod
    def export_docx(text: str, output_path: Path, title: str = None):
        """
        Export as DOCX.

        Args:
            text: Text to export
            output_path: Output file path
            title: Optional document title
        """
        try:
            output_path.parent.mkdir(parents=True, exist_ok=True)

            doc = Document()

            # Add title if provided
            if title:
                doc.add_heading(title, level=1)

            # Add paragraphs
            for para_text in text.split('\n\n'):
                if para_text.strip():
                    p = doc.add_paragraph(para_text.strip())
                    # Set font
                    for run in p.runs:
                        run.font.size = Pt(11)
                        run.font.name = 'Arial'

            doc.save(output_path)
            logger.info(f"Exported DOCX to {output_path}")

        except Exception as e:
            logger.error(f"Failed to export DOCX: {e}")
            raise

    @staticmethod
    def export_all(
        text: str,
        data: Dict,
        base_path: Path,
        formats: List[str] = None,
    ):
        """
        Export to multiple formats.

        Args:
            text: Extracted text
            data: Full OCR data
            base_path: Base output path (without extension)
            formats: List of formats to export
        """
        if formats is None:
            formats = ['txt', 'json']

        for fmt in formats:
            if fmt == 'txt':
                Exporter.export_txt(text, base_path.with_suffix('.txt'))
            elif fmt == 'json':
                Exporter.export_json(data, base_path.with_suffix('.json'))
            elif fmt == 'markdown' or fmt == 'md':
                Exporter.export_markdown(text, base_path.with_suffix('.md'))
            elif fmt == 'docx':
                Exporter.export_docx(text, base_path.with_suffix('.docx'))
            else:
                logger.warning(f"Unknown export format: {fmt}")
